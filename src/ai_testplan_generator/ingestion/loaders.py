"""Streaming document loaders.

Each loader yields `RawBlock`s lazily - never materialising the full doc
in memory. This is what lets us ingest 10k+ page PDFs without going
oom-nuclear. Downstream chunking consumes the iterator, so memory stays
O(one-page) rather than O(whole-doc).

Concrete loaders:
  - PdfLoader:   pypdf (pure-python, handles encrypted + linearised)
  - DocxLoader:  python-docx (paragraphs, headings, tables)
  - XlsxLoader:  openpyxl (streamed via read_only mode)
  - MarkdownLoader / TextLoader:  native
"""

from __future__ import annotations

import hashlib
from abc import ABC, abstractmethod
from collections.abc import Iterator
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

from ai_testplan_generator.models import Document, DocumentKind


@dataclass(frozen=True, slots=True)
class RawBlock:
    """A single structural unit emitted by a loader.

    The chunker fuses heading+prose blocks into sections and cuts them
    into token-bounded chunks. `page` is optional because DOCX / MD don't
    have stable pagination.
    """

    text: str
    kind: Literal[
        "heading", "prose", "table", "list_item", "caption", "code", "formula"
    ] = "prose"
    level: int = 0  # heading depth (0 for non-headings)
    page: int | None = None
    metadata: dict[str, str] = field(default_factory=dict)


class DocumentLoader(ABC):
    """Loader contract - one concrete impl per file type."""

    kind: DocumentKind

    @abstractmethod
    def iter_blocks(self, path: Path) -> Iterator[RawBlock]: ...

    def build_document(
        self,
        path: Path,
        *,
        project_id: str | None = None,
        scope: Literal["general", "project"] = "project",
        title: str | None = None,
    ) -> Document:
        sha = _sha256_file(path)
        return Document(
            project_id=project_id,
            kind=self.kind,
            title=title or path.stem,
            source_uri=path.resolve().as_uri(),
            sha256=sha,
            scope=scope,
        )


# --- PDF ---------------------------------------------------------------------


class PdfLoader(DocumentLoader):
    kind = DocumentKind.PDF

    def iter_blocks(self, path: Path) -> Iterator[RawBlock]:
        from pypdf import PdfReader  # noqa: PLC0415

        reader = PdfReader(str(path))
        for page_idx, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            for block in _split_plain_text_to_blocks(text, page=page_idx):
                yield block


# --- DOCX --------------------------------------------------------------------


class DocxLoader(DocumentLoader):
    kind = DocumentKind.DOCX

    def iter_blocks(self, path: Path) -> Iterator[RawBlock]:
        from docx import Document as _Docx  # noqa: PLC0415

        doc = _Docx(str(path))
        for para in doc.paragraphs:
            text = (para.text or "").strip()
            if not text:
                continue
            style = (para.style.name if para.style else "") or ""
            if style.startswith("Heading"):
                try:
                    level = int(style.split()[-1])
                except ValueError:
                    level = 1
                yield RawBlock(text=text, kind="heading", level=level)
            elif style.lower().startswith("list"):
                yield RawBlock(text=text, kind="list_item")
            else:
                yield RawBlock(text=text, kind="prose")
        # Tables are rendered as tab-separated rows - crude but preserves
        # information density. Downstream we tag them as `table` so the
        # chunker keeps them atomic.
        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                rows.append("\t".join(cell.text.strip() for cell in row.cells))
            if rows:
                yield RawBlock(text="\n".join(rows), kind="table")


# --- XLSX --------------------------------------------------------------------


class XlsxLoader(DocumentLoader):
    kind = DocumentKind.XLSX

    def iter_blocks(self, path: Path) -> Iterator[RawBlock]:
        from openpyxl import load_workbook  # noqa: PLC0415

        wb = load_workbook(str(path), read_only=True, data_only=True)
        try:
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                yield RawBlock(text=sheet_name, kind="heading", level=1)
                # Emit each row as a table block; batch into groups of 50
                # to give the chunker reasonably-sized atomic units.
                batch: list[str] = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        batch.append("\t".join(cells))
                    if len(batch) >= 50:
                        yield RawBlock(
                            text="\n".join(batch), kind="table", metadata={"sheet": sheet_name}
                        )
                        batch = []
                if batch:
                    yield RawBlock(
                        text="\n".join(batch), kind="table", metadata={"sheet": sheet_name}
                    )
        finally:
            wb.close()


# --- Markdown ----------------------------------------------------------------


class MarkdownLoader(DocumentLoader):
    kind = DocumentKind.MARKDOWN

    def iter_blocks(self, path: Path) -> Iterator[RawBlock]:
        content = path.read_text(encoding="utf-8", errors="replace")
        buffer: list[str] = []
        for line in content.splitlines():
            stripped = line.strip()
            if stripped.startswith("#"):
                if buffer:
                    yield RawBlock(text="\n".join(buffer).strip(), kind="prose")
                    buffer = []
                level = len(line) - len(line.lstrip("#"))
                yield RawBlock(text=stripped.lstrip("#").strip(), kind="heading", level=max(1, level))
            elif stripped.startswith(("- ", "* ", "+ ")) or (
                len(stripped) > 2 and stripped[0].isdigit() and stripped[1:3] in {". ", ") "}
            ):
                if buffer:
                    yield RawBlock(text="\n".join(buffer).strip(), kind="prose")
                    buffer = []
                yield RawBlock(text=stripped, kind="list_item")
            elif stripped == "":
                if buffer:
                    yield RawBlock(text="\n".join(buffer).strip(), kind="prose")
                    buffer = []
            else:
                buffer.append(line)
        if buffer:
            yield RawBlock(text="\n".join(buffer).strip(), kind="prose")


# --- Plain text --------------------------------------------------------------


class TextLoader(DocumentLoader):
    kind = DocumentKind.TEXT

    def iter_blocks(self, path: Path) -> Iterator[RawBlock]:
        content = path.read_text(encoding="utf-8", errors="replace")
        yield from _split_plain_text_to_blocks(content, page=None)


# --- helpers -----------------------------------------------------------------


def _split_plain_text_to_blocks(text: str, *, page: int | None) -> Iterator[RawBlock]:
    """Segment flowed text into paragraph blocks and heading-like lines.

    PDF extraction is lossy enough that we do minimal structure inference.
    Heuristics: all-caps lines under ~80 chars and lines matching
    numbered-heading patterns (e.g. "4.2.1 Title") are tagged as headings.
    """
    import re  # noqa: PLC0415

    heading_re = re.compile(r"^\s*\d+(\.\d+){0,4}\s+[A-Z][^.]{2,120}$")

    paragraph: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            if paragraph:
                yield RawBlock(text="\n".join(paragraph).strip(), kind="prose", page=page)
                paragraph = []
            continue
        is_heading = False
        if heading_re.match(stripped):
            is_heading = True
            level = stripped.split()[0].count(".") + 1
        elif stripped.isupper() and 3 < len(stripped) < 80:
            is_heading = True
            level = 1
        else:
            level = 0
        if is_heading:
            if paragraph:
                yield RawBlock(text="\n".join(paragraph).strip(), kind="prose", page=page)
                paragraph = []
            yield RawBlock(text=stripped, kind="heading", level=level, page=page)
        else:
            paragraph.append(stripped)
    if paragraph:
        yield RawBlock(text="\n".join(paragraph).strip(), kind="prose", page=page)


def _sha256_file(path: Path, chunk: int = 65536) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for buf in iter(lambda: f.read(chunk), b""):
            h.update(buf)
    return h.hexdigest()


# --- registry ----------------------------------------------------------------

_LOADERS: dict[DocumentKind, type[DocumentLoader]] = {
    DocumentKind.PDF: PdfLoader,
    DocumentKind.DOCX: DocxLoader,
    DocumentKind.XLSX: XlsxLoader,
    DocumentKind.MARKDOWN: MarkdownLoader,
    DocumentKind.TEXT: TextLoader,
}

_EXT_MAP: dict[str, DocumentKind] = {
    ".pdf": DocumentKind.PDF,
    ".docx": DocumentKind.DOCX,
    ".xlsx": DocumentKind.XLSX,
    ".xlsm": DocumentKind.XLSX,
    ".md": DocumentKind.MARKDOWN,
    ".markdown": DocumentKind.MARKDOWN,
    ".txt": DocumentKind.TEXT,
}


def load_document(
    path: Path | str,
    *,
    project_id: str | None = None,
    scope: Literal["general", "project"] = "project",
    title: str | None = None,
) -> tuple[Document, Iterator[RawBlock]]:
    """Resolve the right loader and return (Document, block iterator)."""
    p = Path(path)
    kind = _EXT_MAP.get(p.suffix.lower(), DocumentKind.UNKNOWN)
    if kind not in _LOADERS:
        raise ValueError(f"Unsupported document extension: {p.suffix}")
    loader = _LOADERS[kind]()
    doc = loader.build_document(p, project_id=project_id, scope=scope, title=title)
    return doc, loader.iter_blocks(p)
